from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


@dataclass(frozen=True)
class ExtractedText:
    text: str
    method: str  # text_extraction | ocr_surya | ocr_easyocr (ocr filled later)
    needs_ocr: bool


def _needs_ocr(page_text: str) -> bool:
    return len(page_text.strip()) < 50


class DocumentProcessor:
    """Extract text from PDF/DOCX and detect OCR need."""

    async def extract(self, file_path: Path, mime: str) -> ExtractedText:
        if mime == "application/pdf":
            return await self._extract_pdf(file_path)
        return await self._extract_docx(file_path)

    async def _extract_pdf(self, file_path: Path) -> ExtractedText:
        doc = fitz.open(file_path)
        parts: list[str] = []
        needs_ocr_any = False
        for page in doc:
            page_text = page.get_text("text") or ""
            if _needs_ocr(page_text):
                needs_ocr_any = True
            parts.append(page_text)
        doc.close()
        text = "\n\n".join(p.strip() for p in parts if p is not None)
        return ExtractedText(text=text.strip(), method="text_extraction", needs_ocr=needs_ocr_any)

    async def _extract_docx(self, file_path: Path) -> ExtractedText:
        d = Document(str(file_path))
        parts: list[str] = []
        for p in d.paragraphs:
            if p.text:
                parts.append(p.text)
        # Tables
        for t in d.tables:
            for row in t.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        return ExtractedText(text=text.strip(), method="text_extraction", needs_ocr=False)

